"""DOCX parser（python-docx）"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from chameleon.api.knowledge.parsers import ParsedDocument


async def parse(source: bytes | str, *, name: str) -> "ParsedDocument":
    from chameleon.api.knowledge.parsers import ParsedDocument

    if isinstance(source, str):
        raise TypeError("docx parser requires bytes input")

    try:
        from docx import Document as DocxDocument
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError("python-docx not installed; uv add python-docx") from e

    doc = DocxDocument(io.BytesIO(source))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 表格内容追加
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs)
    core = doc.core_properties
    metadata: dict = {"name": name}
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author
    return ParsedDocument(text=text, metadata=metadata)
